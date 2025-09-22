import os
from django.conf import settings
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from docx import Document
from datetime import datetime

@csrf_exempt
def guardar_matriz_acuerdos_directiva(request):
    if request.method == "POST":
        try:
            # Crear documento Word
            doc = Document()
            doc.add_heading("Matriz de Acuerdos Directiva", level=1)

            # Crear tabla
            tabla = doc.add_table(rows=1, cols=2)
            hdr_cells = tabla.rows[0].cells
            hdr_cells[0].text = 'Campo'
            hdr_cells[1].text = 'Valor'

            for key, value in request.POST.items():
                row_cells = tabla.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = value

            # Nombre único con fecha/hora
            filename = f"acuerdos_directiva_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            file_path = os.path.join(settings.MEDIA_ROOT, "documentos", filename)

            # Crear carpeta si no existe
            os.makedirs(os.path.dirname(file_path), exist_ok=True)

            # Guardar en el servidor
            doc.save(file_path)

            # Devolver como descarga
            with open(file_path, "rb") as f:
                response = HttpResponse(
                    f.read(),
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                response['Content-Disposition'] = f'attachment; filename="{filename}"'
                return response

        except Exception as e:
            return JsonResponse({"success": False, "error": str(e)})

    return JsonResponse({"success": False, "error": "Método no permitido"})
